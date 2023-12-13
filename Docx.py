from docx import Document

def create_docx(ocr_data, detection_data, docx_path):
    document = Document()
    for page, data in ocr_data.items():
        document.add_heading(f'Page {page}', level=1)
        document.add_paragraph('Extracted Text:')
        document.add_paragraph(' '.join(data['words']))
        document.add_paragraph('Entities:')
        for entity in data['entities']:
            document.add_paragraph(f"{entity['word']}: {entity['label']}", style='List Bullet')
        if int(page) - 1 < len(detection_data):
            document.add_heading('Object Detection:', level=2)
            for detection, b_classes in zip(detection_data[int(page) - 1]['classes'], detection_data[int(page) - 1]['boxes']):
                document.add_paragraph(detection + ':' + str(b_classes), style='List Bullet')
    document.save(docx_path)
